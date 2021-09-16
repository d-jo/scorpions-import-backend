# %%

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

# %%

f = open('./data/test-checkbox.docx', 'rb')
#f = '/home/djo/Documents/UNOMAHA/Fall2021/CSCI4970Capstone/scorpions-academic-import/backend/data/test/word/document.xml'

# import xml.etree.ElementTree as ET
# tree = ET.parse(f)
# root = tree.getroot()

# for child in root:
#   for c2 in child:
#     print(c2.tag, c2.text)
document = Document(f)

# %%

#for p in document.tables:
  #for i in range(0, len(p.columns)):
    #c = p.cell(0, i).text
    #print(c)
    #print(type(c))
  #print('\n')


# %%

#t = document.tables[5]
#for i in range(0, len(t.columns)):
#  c = t.cell(0, i)
#  print(c.text)
#  for cha in c.text:
#    print(cha)
#    print(ord(cha))
# %%

# t = document
# # print(t)
# for i in range(0, len(t.columns)):
#   c = t.cell(3, i)
# #   print(c)
#   c_elm = c._element
#   for cha in c.text:
#     # print(cha)
#     # print(ord(cha))
#     c_text_elm = c._element
#     # print(c_text_elm)
#     checkBoxes = c_text_elm.xpath('.//w:checkBox')
#     for checkBox in checkBoxes:
#       print(checkBox)
#       print('checkBox value is %s' % checkBox.get(qn('w:val')))


# doc_elm = document._element
# print(doc_elm)
# checkBoxes = doc_elm.xpath('.//w14:checkBox')
# print(checkBoxes)
# for checkBox in checkBoxes:
#     print('checkBox value is %s' % checkBox.get(qn('w:val')))


# for paragraph in document.paragraphs:
#     if paragraph.text == '':
#         continue
#     # print(paragraph.text)
#     p = paragraph._element
#     # print(p)
#     checkBoxes = p.xpath('//w:p')
#     temp = p.xpath('//w14:checkbox')
#     if temp:
#         for boxes in temp:
#             print(temp[0].text)
#             # print(paragraph.text)
#             # print(p.xml)
#             # something = boxes.xpath('//w14:checked')
#             # print(something)
#             print('checkBox value is %s' % boxes.get(qn('w14:val')))

#     if checkBoxes:
#         # print(checkBoxes)
#         # print(paragraph.text)
#         # print(p.xml)
#         print('')


for paragraph in document.paragraphs:
    if paragraph.text == '':
        continue
    print(paragraph)
    checkboxes = etree.ElementBase.xpath(paragraph._element, './/w14:checkbox', namespaces=paragraph._element.nsmap)
    for checkbox in checkboxes:
        # print(paragraph.text)
        p = paragraph._element
        # print(p.xml)
        # print(checkbox)
        for item in checkbox.findall('.//w14:checkbox', namespaces=paragraph._element.nsmap):
            print(item)
            print(etree.dump(item))
        # something = tree.get("w14:checked")
        # print(something)
        # print(checkbox)
        # print('checkBox value is %s' % checkbox.get('w14:checked'))



# %%
