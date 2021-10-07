# %%
from docx import Document
from lxml import etree
import xml.etree.ElementTree as ET


def rec_traverse(tr):
    result = []
    if tr.tag.endswith("t"):
        result.append((tr.tag, tr.text))
        pass
    if tr.tag.endswith("checked"):
        for attr in tr.attrib:
            if attr.endswith("val"):
                checked = tr.attrib[attr]
                result.append(("checkbox", checked))
                
    #if tr.tag.endswith()
    for c in tr:
        result.extend(rec_traverse(c))
    
    return result


def is_checked(docElements):
    """Not able to detemine if checkbox is checked or not."""
    for docElem in docElements:
        checkboxes = etree.ElementBase.xpath(docElem._element, './/w14:checkbox', namespaces=docElem._element.nsmap)
        # for checkbox in checkboxes:
        #     print(checkbox)
        p = docElem._element
        tree=ET.fromstring(p.xml)
        # print(tree)
        for a in rec_traverse(tree):
            print('****************************************************\n')
            print(a)
            # print(checkbox)
            # print(p.xml)
            # for item in checkbox.findall('.//w14:checkbox', namespaces=docElem._element.nsmap):
                # print(item)
                # print(etree.dump(item))


def read_document(document):
    """Read in document and print out the paragraph contents."""
    documentText = []
    documentElements = []
    for paragraph in document.paragraphs:
        if paragraph.text == '':
            continue
        documentText.append(paragraph.text)
        documentElements.append(paragraph)
        # print(paragraph)
    tablesCells = []
    tablesCellsText = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text == '':
                        continue
                    tablesCells.append(para)
                    tablesCellsText.append(para.text.strip())
    print('Document plain text\n')
    print('****************************************************\n')
    print('\n'.join(documentText))
    print('****************************************************\n')


    print('Document table text\n')
    print('****************************************************\n')
    print('\n'.join(tablesCellsText))
    print('****************************************************\n')
    # is_checked(documentElements)
    is_checked(tablesCells)


f = open('./data/grad2018-regular.docx', 'rb')
document = Document(f)
read_document(document)
# %%

# w:tc
#   w:p
#       w:sdt -> w:sdtPr
#           w14:checkbox -> w14:checked
#       w:r
#           w:t = text
#           
#      

# %%

import xml.etree.ElementTree as ET


tree=ET.parse('data/test/word/document.xml')


# %%
