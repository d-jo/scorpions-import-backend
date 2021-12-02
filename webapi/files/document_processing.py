# %%
from docx import Document
import re
from difflib import SequenceMatcher
from models.model import *

# %%
def process_report(filename):
    """
    Process report for data extraction

    param filename: filename to open and extract data from
    """
    return read_document(Document(open(filename, 'rb')))

def read_document(document):
    """
    Read in document and extract all data from it.
    Surrounds each data type with try catches to not completely halt execution,
    some data parsed is better than none!
    :param document: opened document from python-docx.
    """
    report = get_report_info(document)
    
    slos = []
    measures = [[]]
    analysisList = [[]]
    decisions = []
    methods = []
    adaList = []
    for table in get_table_cells(document):
        # Empty table
        if len(table) < 1: 
            continue
        # SLO data
        if SequenceMatcher(None, table[0]['cell'], "Student Learning Outcomes").ratio() >= 0.9:
            try:
                slos = get_slo_data(table)
            except:
                print("Error parsing slo data!")
        # Assessmment data
        elif re.match(re.compile('SLO..:'), table[0]['cell']):
            # measures.append(get_measures_data(table))
            #TODO figure out how to make this work in a method call
            try:
                sloNum = ""
                state = ""
                measure = Measure().to_dict()
                for cell in table:
                    if state != "":
                        measure = state_match(state, measure, cell)
                        state = ""
                        continue
                    # if we have a new SLO num, add it to the list and move to the next
                    # Note: there may be multiple measures for one SLO, need to preserve number
                    if re.match(re.compile('^SLO..:'), cell['cell']):
                        num = extract_text(cell['cell'], "SLO ")[0]
                        if sloNum != "":
                            idx = int(sloNum)
                            while len(measures) < idx:
                                measures.append([]) 
                            measureobj = Measure()
                            measureobj.init_from_dict(measure)
                            measures[idx-1].append(measureobj)
                            measure = Measure().to_dict()
                        sloNum = num
                        measure['slo_id'] = sloNum
                        continue
                    if "Title" in cell['cell']:
                        measure['title'] = extract_text(cell['cell'], ":").lstrip()
                        continue
                    if "Measure Aligns to the SLO" in cell['cell']:
                        measure['description'] = extract_text(cell['cell'], "SLO").lstrip()
                        continue
                    for st in ["Domain", "Type", "Point in Program", "Population", "Frequency", "Threshold", "Program"]:
                        if st in cell['cell']:
                            state = map_state(st)
                            break
                idx = int(sloNum)
                while len(measures) < idx:
                    measures.append([]) 
                measureobj = Measure()
                measureobj.init_from_dict(measure)
                measures[idx-1].append(measureobj)
            except:
                print("Error parsing measure!")
        # Analysis Data
        elif is_analysis(table):
            try:
                analysisList = get_analysis_data(table)
            except:
                print("Error parsing analysis data!")
        # Decisions/Actions Data
        elif re.match(re.compile('^SLO..'), table[0]['cell']) and not status_table(table):
            try:
                decisions = get_decisions_data(table)
            except:
                print("Error parsing decisions data!")
        elif is_methods(table):
            try:
                methods = get_methods_data(table)
            except:
                print("Error parsing methods data!")
        elif is_acc_data_analysis(table):
            try:
                adaList = get_acc_data_analysis(table)
            except:
                print("Error parsing accredited data analysis!")
        else:
            print("Misc data found")
    
    return [report, slos, measures, analysisList, decisions, methods, adaList]

#TODO clean up, possibly make a process_engine class with below methods to leave this file cleaner

def get_acc_data_analysis(table):
    """
    parses the given table to extract ada data and return a list of ada
    :param table: the table to get ada data from
    """
    sloNum = ""
    adaList = []
    ada = AccreditedDataAnalysis()
    for cell in table:
        if "SLO " in cell['cell']:
            for c in cell['cell']:
                if c.isdigit():
                    num = c
                    break
            if sloNum != "":
                adaList.append(ada)
                ada = AccreditedDataAnalysis()
            sloNum = num
            ada.slo_id = sloNum
        elif not cell['checkboxes']:
            text = cell['cell']
            for c in text:
                if isinstance(c, str) and is_checkbox(c):
                    parts = text.split(chr(9746))  
                    if len(parts) > 1:
                        for p in parts:
                            ada.status = get_first_word(p).strip() 
                    break
        elif cell['checkboxes']:
            pos = 0
            for cb in cell['checkboxes']:
                for child in cb.getchildren():
                    if child.tag.endswith("checked"):
                        if(int(re.search(r'\d+', child.values()[0]).group())):
                            text = get_word_at(pos, cell['cell'])
                            if text != "" and "Met" not in text and text != "Unknown":
                                text += " Met" 
                            ada.status = text
                        pos += 1
    if ada not in adaList: adaList.append(ada)
    return adaList

def is_acc_data_analysis(table):
    """
    checks if the table is likely an ada table to extract data from

    :param table: the table to check
    """
    for cell in table:
        if "Met" in cell['cell'] and "Partially Met" in cell['cell'] and "Not Met" in cell['cell'] and "Unknown" in cell['cell']:
            return True
    return False

def get_methods_data(table):
    """
    parses the given table to extract method data and return a list of methods
    :param table: the table to get method data from
    """
    sloNum = ""
    methodList = []
    method = Methods()
    for cell in table:
        if cell['cell'].isdigit():
            if sloNum != "":
                methodList.append(method)
                method = Methods()
            sloNum = cell['cell']
            method.slo_id = cell['cell']
        elif method.measure == "":
            method.measure = cell['cell']
        elif method.domain == "":
            method.domain = cell['cell']
        else : 
            method.data_collection = cell['cell']
    if method not in methodList: methodList.append(method)
    return methodList

#Try to get this working with the Assessment data in read_document
def get_measures_data(table):
    """
    ***NOTE: This currently will not work see read_document under Measures branch to see implementation***
    parses the given table to extract measures data and return a list of measures
    :param table: the table to get measure data from
    """
    measures = []
    for cell in table:
        print(cell)
    return measures

def get_analysis_data(table):
    """
    parses the given table to extract analysis data and return a list of analysis
    :param table: the table to get analysis data from
    """
    sloNum = ""
    analysisList = [[]]
    analysis = CollectionAnalysis()
    for cell in table:
        # if we have a new SLO num, add it to the list and move to the next
        # Note: there may be multiple analysis for one SLO, need to preserve number
        if "SLO " in cell['cell']:
            for c in cell['cell']:
                if c.isdigit():
                    num = c
                    break
            if sloNum != "":
                idx = int(sloNum)
                while len(analysisList) < idx:
                    analysisList.append([])
                analysisList[idx-1].append(analysis)
                analysis = CollectionAnalysis()
            sloNum = num
            analysis.slo_id = sloNum
        elif "%" in cell['cell']:
            analysis.percentage_who_met_or_exceeded = cell['cell']
        elif is_year(cell['cell']):
            analysis.data_collection_date_range = cell['cell']
        elif not re.match(re.compile('^Note:'), cell['cell']):
            analysis.number_of_students_assessed = cell['cell']
    idx = int(sloNum)
    while len(analysisList) < idx:
        analysisList.append([])
    analysisList[idx-1].append(analysis)
    return analysisList

def get_decisions_data(table):
    """
    parses the given table to extract decisions data and return a list of decisions
    :param table: the table to get decisions data from
    """
    sloNum = ""
    decisions = []
    decision = DecisionsAction()
    for cell in table:
        # if we have a new SLO num, add it to the list and move to the next
        if re.match(re.compile('^SLO..$'), cell['cell']):
            num = cell['cell'].split("SLO ")[1]
            if sloNum != "":
                decisions.append(decision)
                decision = DecisionsAction()
            sloNum = num
            decision.slo_id = sloNum
        else:
            decision.content = cell['cell']
    decisions.append(decision)
    return decisions

def status_table(table):
    """
    checks if the table is likely a status table to extract data from

    :param table: the table to check
    """
    for cell in table:
        if "Partially Met" in cell['cell']:
            return True
    return False
 
def is_year(text):
    """
    checks if the text is a year, used to differentiate data fields
    very likely the forms will not have anything before 1900s so if it is above 1900, its a year 
    :param text: the text to check
    """
    year = ""
    found = False
    for c in text:
        if c.isdigit():
            year += c
            found = True
        elif found:
            break
    if year != "":
        return int(year) > 1900
    return False

def is_analysis(table):
    """
    checks if the table is likely an analysis table to extract data from

    :param table: the table to check
    """
    for cell in table:
        if SequenceMatcher(None, cell['cell'], "Data Collection Date Range").ratio() >= 0.95:
            return True
    return False

def is_methods(table):
    """
    checks if the table is likely a methods table to extract data from

    :param table: the table to check
    """
    for cell in table:
        if ("Product" in cell['cell'] and "Performance" in cell['cell'] and "Examination" in cell['cell']):
            return True
    return False

def get_slo_data(table):
    """
    parses the given table to extract slo data and return a list of slos
    :param table: the table to get slo data from
    """
    slos = []
    slo = SLO()
    for cells in table:
        slo = slo_matcher(cells['cell'], slo)
        if(slo.description != ""):
            if not has_duplicate(slos, slo):
                slos.append(slo)
            slo = SLO()
        sloNum = -1
        if slo.id != "":
            parts = slo.id.split("SLO")
            if len(parts) > 1:
                sloNum = int(parts[1].lstrip())
        if sloNum == -1:
            sloNum = max(0, len(slos) - 1)
        if not cells['checkboxes']:
            text = cells['cell']
            for c in text:
                if isinstance(c, str) and is_checkbox(c):
                    parts = text.split(chr(9746))  
                    if len(parts) > 1:
                        for p in parts:
                            result = get_first_word(p).strip()
                            slos = slo_attr_match(result, slos, sloNum)
                    break
        elif cells['checkboxes']:
            pos = 0
            for cb in cells['checkboxes']:
                for child in cb.getchildren():
                    if child.tag.endswith("checked"):
                        if(int(re.search(r'\d+', child.values()[0]).group())):
                            word = get_word_at(pos, cells['cell'])
                            slos = slo_attr_match(word, slos, sloNum)
                        pos += 1
    return slos

# TODO eventually instead of saving all data in a list(memory heavyish), just extract data here?
def get_table_cells(document):
    """
    prepare table cells to be parsed through. This organizes our data in separate elements that we can loop over.
    Nore: we hog up more running memory because we save an array of table 
    and have the file open with the table data still in it. 
    The files aren't too large so this probably isn't a big issue

    :param document: the document to get tables from
    """
    tableCells = [[]]
    for table in document.tables:
        cell_information = []
        for row in table.rows:
            # cells get duplicated if they are merged cells
            for cell in iter_unique_cells(row):
                p = cell._element
                checkboxes = p.xpath('.//w14:checkbox')
                val = {"cell":  cell.text.strip(), "checkboxes": checkboxes}
                cell_information.append(val)
        tableCells.append(cell_information)
    return tableCells

def get_report_info(document):
    """
    this method gets the main report info from the document, 
    usually all the header data. Parses through all the text in the file
    so early exit if all data is found is placed here

    :param document: the document to get report info from
    """
    report = Report()
    for paragraph in document.paragraphs:
        if paragraph.text == '':
            continue
        if report.title == "":
            report.title = paragraph.text
        report = report_matcher(paragraph.text, report)
        if is_full(report): #early exit for faster time but can be removed
            return report
    return report
    
def map_state(st):
    """
    maps a state to a Measure attribute, this makes 
    setting an attribute with a dict very easy

    :param st: state to map
    """
    if st == "Domain":
        return "domain"
    if st == "Type":
        return "type"
    if st == "Point in Program":
        return "point_in_program"
    if st == "Population":
        return "population_measured"
    if st == "Frequency":
        return "frequency_of_collection"
    if st == "Threshold":
        return "proficiency_threshold"
    if st == "Program":
        return "proficiency_target"

def state_match(state, m, cell):
    """
    Given a state, find the necessary data in the cell
    and then set that in the measure object

    :param state: state used to set attribute of Measure
    :param m: Measure to update
    :param cell: cell data to look through
    """
    if state == "proficiency_threshold" or state == "proficiency_target":
        if cell['cell'] != "Describe:":
            m[state] = cell['cell']
        return m
    checked = []
    if cell['checkboxes']:
        pos = 0
        for cb in cell['checkboxes']:
            for child in cb.getchildren():
                if child.tag.endswith("checked"):
                    if(int(re.search(r'\d+', child.values()[0]).group())):
                        words = re.findall('[a-zA-Z][^A-Z]*', cell['cell'])
                        if "Sample of students" in words[pos]:
                            idx = 0
                            text = ""
                            for items in words:
                                if idx >= pos:
                                    text += words[idx]
                                idx += 1
                            checked.append(text.lstrip())
                        elif "Direct" in words[pos]:
                            checked.append(words[pos] + " " + words[pos+1])
                        else:
                            checked.append(words[pos].lstrip())
                    pos += 1
    else:
        text = cell['cell']
        items = text.split(chr(9746))
        items.pop(0)
        for text in items:
            checked.append(text.split(chr(9744))[0].lstrip())
    if len(checked) > 0:
        m[state] = ', '.join(checked)
    return m

def iter_unique_cells(row):
    """
    Generate cells in `row` skipping empty grid cells.
    Without this, there will be duplicate cells in table cells

    :param row: row to have dupes removed
    """
    prior_tc = None
    for cell in row.cells:
        this_tc = cell._tc
        if this_tc is prior_tc:
            continue
        prior_tc = this_tc
        yield cell

def get_word_at(pos, text):
    """
    Breaks up the sentence into individual words 
    and gives the word at the specified position

    :param text: text to break up
    :param pos: position to get the word at
    """
    words = []
    inword = 0
    for c in text:
        if c in " \r\n\t":
            inword = 0
        elif not inword:
            words = words + [c]
            inword = 1
        else:
            words[-1] = words[-1] + c
    return words[pos]

def slo_attr_match(word, slos, sloNum):
    """
    Match a word to an slo attribute for setting data

    :param word: word to match to slo attribue
    :param slo: slo to set data into
    """

    # rework this as taxonomy levels may not be all of these
    # if the field is empty, first entry doesn't need a comma, else add a comma to separate
    if(word in ["Knowledge", "Analysis", "Comprehension","Synthesis","Application", "Evaluation"]):
        slos[sloNum].bloom += word if slos[sloNum].bloom == "" else ", " + word  
    elif(word in ["1", "2", "3", "4", "Not applicable for SLO"]):
        slos[sloNum].common_graduate_program_slo += word if slos[sloNum].common_graduate_program_slo == "" else ", " + word
    return slos

def get_first_word(str):
    """
    Returns the first word in a string

    :param str: string to retrieve word from
    """
    result = ""
    special_chars = "/ \\@!&*().?,"
    for c in str:
        if is_checkbox(c):
            return result
        elif c.isalpha() or c in special_chars:
            result += c
    return result

def is_checkbox(c):
    """
    checks if given character is a ascii checkbox
 
    :param c: char to check
    """
    number = ord(c)
    return not c.isalpha() and number == 9744 or number == 9746

def has_duplicate(slos, slo):
    """
    Take a list of slos and determine if it already has the same slo in it

    :param slos: list of slos.
    :param slo: slo to be checked against.
    """
    for sloItem in slos: 
        if SequenceMatcher(None, sloItem.description, slo.description).ratio() >= 0.95:
            return True    
    return False

def is_full(report):
    """
    checks if report data is filled to exit early
 
    :param report: report to check
    """
    return (report.college != "" and report.department != "" 
            and report.program != "" and report.degree_level != "" 
            and report.academic_year != "" and report.date_range != "" 
            and report.author != "")

def report_matcher(str, report:Report):
    """
    Reads a string to extract the main report data

    :param str: string to look at for data.
    :param report: report object that will be updated and returned.
    """
    if report is None:
        report = Report()
    if "College:" in str:
        report.college = extract_text(str, "College:")
    if "Department/School:" in str:
        report.department = extract_text(str, "Department/School:")
    if "Program:" in str:
        report.program = extract_text(str, "Program:")
    if "Degree Level:" in str:
        report.degree_level = extract_text(str, "Degree Level:")
    if "Academic Year of Report:" in str:
        report.academic_year = extract_text(str, "Academic Year of Report:")
    if "Date Range of Reported Data:" in str:
        report.date_range = extract_text(str, "Date Range of Reported Data:")
    if "Person Preparing the Report:" in str:
        report.author = extract_text(str, "Person Preparing the Report:")
    return report

def slo_matcher(str, slo:SLO):
    """
    Reads a string to extract the main slo data

    :param str: string to look at for data.
    :param slo: slo object that will be updated and returned.
    """
    if slo is None:
        slo = SLO()
    regex = re.compile('^SLO..:')
    if re.match(regex, str):
        for text in re.split(":", str):
            if "SLO" in text:
                slo.id = text
            elif(text != ""):
                slo.description = text.strip()
    elif slo.description == "" and contains_one_number(str):
        for c in str:
            if c.isdigit():
                slo.id = c
                slo.description = str.split(c)[1]
    return slo

def contains_one_number(str):
    """ 
    checks if the given string has only one number in it
    this helps differentiate between an slo number with a description - returns true
    or if it is the common graduate data - returns false
    :param str: the string to check.
    """
    found = False
    for c in str:
        if(c.isdigit() and found): return False
        elif(c.isdigit()): found = True
    return True

def extract_text(str, split_point):
    """
    Reads a string to split text and remove tab characters.
    Some text separate information with tabs so we only want the beginning
    ex:
        College:\\t<text>\\t\\t\\t <other text>
    we want just <text> so split at College: and \\t

    :param str: string to be split.
    :param split_point: string match to split at.
    """
    pieces = str.split(split_point)
    if(len(pieces) < 2):
        return ""
    sections = pieces[1].split("\t")
    for parts in sections:
        toReturn = parts.lstrip() 
        if toReturn != '':
            return toReturn
    return ""    